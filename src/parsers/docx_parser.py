"""DOCX resume parsing."""


def parse_docx(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Install python-docx to parse DOCX files.") from exc

    document = Document(uploaded_file)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()
